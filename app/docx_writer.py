from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold =True
    run.underline = True
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.font.size = Pt(14)

def add_subheading(doc, text):
    paragraph = doc.add_paragraph()
    run =  paragraph.add_run(text)
    run.bold =  True
    run.font.size = Pt(12)

def add_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        start, end = match.span()
        paragraph.add_run(text[cursor:start])
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True
        cursor = end

    paragraph.add_run(text[cursor:])

def add_code_block(doc, code_lines):
    para =doc.add_paragraph()
    run =para.add_run("\n".join(code_lines))
    # run.font.name ="Courier New"
    run.font.size = Pt(12)

# def add_markdown_table(doc, lines):
#     headers = [cell.strip(" *") for cell in lines[0].split("|") if cell.strip()]
#     num_cols = len(headers)

#     table = doc.add_table(rows=1, cols=num_cols)
#     table.style = 'Table Grid'

#     hdr_cells = table.rows[0].cells
#     for i, h in enumerate(headers):
#         hdr_cells[i].text = h

#     for line in lines[2:]:  # skip header and separator
#         cells = [cell.strip() for cell in line.split("|") if cell.strip()]

#         # ⚠️ Align row cells to number of columns
#         if len(cells) < num_cols:
#             cells += [""] * (num_cols - len(cells))  # pad short rows
#         elif len(cells) > num_cols:
#             cells = cells[:num_cols]  # truncate long rows

#         row_cells = table.add_row().cells
#         for i, cell in enumerate(cells):
#             row_cells[i].text = cell
def add_markdown_table(doc, lines):
# """Improved table creation with better error handling"""
    if not lines or len(lines) < 2:
        print("Warning: Invalid table data - need at least header and separator")
        return
    
    # Extract headers - handle both | cell | and |cell| formats
    header_line = lines[0].strip()
    if header_line.startswith('|') and header_line.endswith('|'):
        header_line = header_line[1:-1]  # Remove outer pipes
    
    headers = [cell.strip() for cell in header_line.split("|") if cell.strip()]
    
    if not headers:
        print("Warning: No valid headers found in table")
        return
    
    num_cols = len(headers)
    print(f"Creating table with {num_cols} columns: {headers}")
    
    # Create table
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Make headers bold
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Process data rows (skip header and separator line)
    data_lines = lines[2:] if len(lines) > 2 else []
    
    for line_num, line in enumerate(data_lines):
        line = line.strip()
        if not line:
            continue
            
        # Remove outer pipes if present
        if line.startswith('|') and line.endswith('|'):
            line = line[1:-1]
        
        # Split cells
        raw_cells = line.split("|")
        cells = [cell.strip() for cell in raw_cells]
        
        # Handle empty or short rows
        if not any(cells):  # Skip completely empty rows
            continue
            
        # Pad or truncate to match column count
        if len(cells) < num_cols:
            cells.extend([""] * (num_cols - len(cells)))
        elif len(cells) > num_cols:
            cells = cells[:num_cols]
        
        # Add row to table
        row_cells = table.add_row().cells
        for i, cell_content in enumerate(cells):
            row_cells[i].text = cell_content
    
    print(f"Table created successfully with {len(table.rows)} rows")

def is_table_line(line):
    """Better table line detection"""
    line = line.strip()
    if not line:
        return False
    
    # Check for pipe characters (markdown table indicator)
    if '|' in line:
        # Must have at least 2 pipes or be enclosed in pipes
        pipe_count = line.count('|')
        if pipe_count >= 2:
            return True
        if line.startswith('|') and line.endswith('|') and pipe_count >= 2:
            return True
    
    return False

def is_table_separator(line):
    """Detect table separator line (e.g., |---|---|)"""
    line = line.strip()
    if not line:
        return False
    
    # Remove outer pipes
    if line.startswith('|') and line.endswith('|'):
        line = line[1:-1]
    
    # Check if line contains only dashes, pipes, and colons (for alignment)
    import re
    return bool(re.match(r'^[\s\-\|\:]+$', line)) and '-' in line

# def add_markdown_table(doc, lines):
#     headers = [cell.strip(" *") for cell in lines[0].split("|") if cell.strip()] 
#     rows = [
#         [cell.strip() for cell in row.split("|") if cell.strip()]
#         for row in lines[2:]# skip header and separator
#     ]
#     table =  doc.add_table(rows=1, cols=len(headers))
#     table.style = 'Table Grid'
#     hdr_cells=  table.rows[0].cells
#     for i, h in enumerate(headers):
#         hdr_cells[i].text = h
#     for row in rows:
#         row_cells= table.add_row().cells 
#         for i, cell in enumerate(row):
#             row_cells[i].text = cell

def create_docx(ts_text: str, buffer):
    doc = Document() 
    doc.add_heading('TECHNICAL SPECIFICATION', level=1)

    lines = ts_text.splitlines()
    current_section = ""
    current_content = []
    in_code_block = False
    code_block_lines = []
    in_table =False
    table_lines = []

    #section_header_pattern= re.compile(r"^\s*(\d{1,2})\.\s"(.?):\s(.+)?$")
    #plain_header_pattern = re.compile(r"^\s*(\d{1,2})\. \s*(.+)$")
    #subheading pattern = re.compile(r"^\s*(\d{1,2})\.(\d+)\s*(.+?):25")

    table_line_pattern = re.compile(r"^\|(.+?)\|$")
    section_heading_pattern = re.compile(r"^\d+(\.\d+)*\s.*")
    def flush_current_content():
        if current_section:
            add_heading(doc, current_section)
        for para in current_content:
            add_paragraph(doc, para)

    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Handle code block
        if line.startswith("```"):
            in_code_block = not in_code_block
            if not in_code_block:
                add_code_block(doc, code_block_lines)
                code_block_lines = []
            continue
        elif in_code_block:
            code_block_lines.append(line)
            continue
        #Handle markdown table
        if table_line_pattern.match(line):
            table_lines.append(line)
            in_table =  True
            continue
        elif in_table and not table_line_pattern.match(line):
            flush_current_content()
            current_content = []
            add_markdown_table(doc, table_lines)
            table_lines =[]
            in_table= False
            continue
        #Section header with content (e.g., "1. Title: My Program")
        # match_full section_header_pattern.match(line)
        # match plain plain_header_pattern.match(line)
        #match_sub subheading pattern.match(line)

        #if match full:
        if current_section and current_content:
            add_heading(doc, current_section)
            for content in current_content:
                add_paragraph(doc, content)
            current_content = []

            #current_section f"{match_full.group(1)). (match_full.group(2))"
            #current_content [match_full.group(3)] if match_full.group(3) else []
        
        #elif match plain:
        if current_section and current_content:
            add_heading(doc, current_section)
            for content in current_content:
                add_paragraph(doc, content)
            current_content = []

            #current section f"(match_plain.group(1)). (match_plain.group(2)}"
            #current_content = []

        # elif match sub:
        if current_content:
            for content in current_content:
                add_paragraph(doc, content)
            current_content = []

            #subheader-f(match_sub.group(1)).(match_sub.group(2)) (match_sub.group(3))"
            #add_subheading(doc, subheader)
        else:
            current_content.append(line)
    # Final flush
    if current_section and current_content:
        add_heading(doc, current_section)
        for content in current_content:
            add_paragraph(doc, content)
    doc.save(buffer)